"""Utilities for DOCX parsing and content replacement using python-docx."""
import copy
import re
from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Section extraction
# ---------------------------------------------------------------------------

def _is_heading_style(paragraph) -> bool:
    """Return True if the paragraph uses a Heading 1 or Heading 2 style."""
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.startswith("Heading 1") or style_name.startswith("Heading 2")


def _is_allcaps_bold(paragraph) -> bool:
    """Return True if the paragraph looks like an ALL-CAPS bold section header."""
    text = paragraph.text.strip()
    if not text:
        return False
    # Must be entirely uppercase (letters + common punctuation)
    if not re.match(r'^[A-Z][A-Z0-9 /&()\-:]+$', text):
        return False
    # At least one run must be bold, or all text runs contain bold
    runs = paragraph.runs
    if not runs:
        return False
    return any(r.bold for r in runs if r.text.strip())


def extract_sections(docx_bytes: bytes) -> list[dict]:
    """
    Parse a DOCX and return a list of sections.

    Each section dict: {"heading": str, "paragraphs": list[str]}

    Uses Heading 1 / Heading 2 styles to detect section boundaries.
    Falls back to ALL-CAPS bold paragraphs if no heading styles are found.
    Returns all content, including a synthetic "header" section for content
    before the first heading (name, contact info).
    """
    doc = Document(BytesIO(docx_bytes))
    paragraphs = doc.paragraphs

    # Determine detection strategy
    has_heading_styles = any(_is_heading_style(p) for p in paragraphs)

    def _is_section_boundary(p) -> bool:
        if has_heading_styles:
            return _is_heading_style(p)
        return _is_allcaps_bold(p)

    sections: list[dict] = []
    current_heading: Optional[str] = None
    current_paragraphs: list[str] = []

    for para in paragraphs:
        text = para.text.strip()

        if _is_section_boundary(para):
            # Save the previous section
            if current_heading is not None or current_paragraphs:
                sections.append({
                    "heading": current_heading or "header",
                    "paragraphs": current_paragraphs,
                })
            current_heading = text
            current_paragraphs = []
        else:
            # Only include non-empty lines, but preserve blank lines as empty strings
            # so downstream consumers know about spacing
            current_paragraphs.append(text)

    # Flush the last section
    if current_heading is not None or current_paragraphs:
        sections.append({
            "heading": current_heading or "header",
            "paragraphs": current_paragraphs,
        })

    # If nothing was detected, return everything as a single header section
    if not sections:
        all_text = [p.text.strip() for p in paragraphs]
        sections = [{"heading": "header", "paragraphs": all_text}]

    return sections


# ---------------------------------------------------------------------------
# Section application (content replacement preserving styles)
# ---------------------------------------------------------------------------

def _clear_paragraph_text(paragraph) -> None:
    """Remove all text from a paragraph while keeping the first run's formatting."""
    for run in paragraph.runs:
        run.text = ""


def _clone_paragraph_after(target_paragraph, reference_paragraph) -> object:
    """
    Insert a new paragraph after *target_paragraph*, cloned from
    *reference_paragraph* (copies XML element including style).
    Returns the new paragraph object.
    """
    new_elem = copy.deepcopy(reference_paragraph._element)
    # Clear all run text in the clone
    for r in new_elem.findall(qn("w:r")):
        for t in r.findall(qn("w:t")):
            t.text = ""
    target_paragraph._element.addnext(new_elem)

    # Wrap as a Paragraph object from the same document
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.oxml import OxmlElement  # noqa: F401
    return DocxParagraph(new_elem, target_paragraph._parent)


def _set_paragraph_text(paragraph, text: str) -> None:
    """
    Set the paragraph text, preserving the formatting of the first run.
    If there are multiple runs, collapse them into a single run.
    """
    runs = paragraph.runs
    if runs:
        # Keep the first run's formatting, set its text, clear the rest
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        # No runs — add one
        paragraph.add_run(text)


def _find_section_paragraphs(doc: Document, heading: str) -> tuple[int, int]:
    """
    Return (heading_index, first_body_index) where heading_index is the index
    of the paragraph matching *heading* and first_body_index is the next paragraph.
    Returns (-1, -1) if not found.
    """
    paragraphs = doc.paragraphs
    heading_lower = heading.strip().lower()
    for i, para in enumerate(paragraphs):
        if para.text.strip().lower() == heading_lower:
            return i, i + 1
    return -1, -1


def _find_section_body_range(doc: Document, heading: str) -> tuple[int, int]:
    """
    Return (body_start, body_end) indices into doc.paragraphs for paragraphs
    belonging to *heading*'s section (i.e., between this heading and the next).
    Returns (-1, -1) if heading not found.
    """
    paragraphs = doc.paragraphs
    heading_lower = heading.strip().lower()

    # Determine which detection strategy was used in this document
    has_heading_styles = any(_is_heading_style(p) for p in paragraphs)

    def _is_boundary(p) -> bool:
        if has_heading_styles:
            return _is_heading_style(p)
        return _is_allcaps_bold(p)

    heading_idx = -1
    for i, para in enumerate(paragraphs):
        if para.text.strip().lower() == heading_lower:
            heading_idx = i
            break

    if heading_idx == -1:
        return -1, -1

    body_start = heading_idx + 1
    body_end = len(paragraphs)
    for i in range(body_start, len(paragraphs)):
        if _is_boundary(paragraphs[i]) and paragraphs[i].text.strip():
            body_end = i
            break

    return body_start, body_end


def apply_sections_to_docx(source_bytes: bytes, sections: list[dict]) -> bytes:
    """
    Given the original DOCX bytes and a list of {heading, content: str} dicts
    from the LLM, produce a new DOCX with new content while preserving paragraph
    styles from the source.

    Strategy:
    - Copy the source document in memory
    - For each section in the LLM output, find the corresponding heading in the doc
    - Replace body paragraph text while preserving Run formatting
    - If the LLM section has more paragraphs than the original, clone the last
      paragraph's style for the extras
    - If fewer paragraphs, trim extras (remove their text, leave as empty paragraphs
      or remove them entirely)
    - Return as bytes
    """
    doc = Document(BytesIO(source_bytes))

    for section in sections:
        heading: str = section.get("heading", "")
        content: str = section.get("content", "")

        # Split content into lines; filter empty-only lines at start/end
        new_lines = content.split("\n")
        # Strip leading/trailing empty lines
        while new_lines and not new_lines[0].strip():
            new_lines.pop(0)
        while new_lines and not new_lines[-1].strip():
            new_lines.pop()

        if heading.lower() == "header":
            # Special case: update lines before the first real section heading
            has_heading_styles = any(_is_heading_style(p) for p in doc.paragraphs)

            def _is_boundary(p) -> bool:
                if has_heading_styles:
                    return _is_heading_style(p)
                return _is_allcaps_bold(p)

            header_paragraphs = []
            for p in doc.paragraphs:
                if _is_boundary(p) and p.text.strip():
                    break
                header_paragraphs.append(p)

            _replace_paragraph_range(header_paragraphs, new_lines)
            continue

        body_start, body_end = _find_section_body_range(doc, heading)
        if body_start == -1:
            # Heading not found — skip this section
            continue

        body_paragraphs = doc.paragraphs[body_start:body_end]
        _replace_paragraph_range(body_paragraphs, new_lines)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _replace_paragraph_range(paragraphs: list, new_lines: list[str]) -> None:
    """
    Replace the text of *paragraphs* with *new_lines*, preserving formatting.
    Clones the last paragraph style for extra lines; clears extra paragraphs
    if new_lines is shorter.
    """
    if not paragraphs:
        return

    n_orig = len(paragraphs)
    n_new = len(new_lines)

    # Update existing paragraphs
    for i in range(min(n_orig, n_new)):
        _set_paragraph_text(paragraphs[i], new_lines[i])

    if n_new > n_orig:
        # Need to insert extra paragraphs, cloning the last original paragraph
        reference = paragraphs[-1]
        last = reference
        for line in new_lines[n_orig:]:
            new_para = _clone_paragraph_after(last, reference)
            _set_paragraph_text(new_para, line)
            last = new_para
    elif n_new < n_orig:
        # Clear extra paragraphs (set text to empty string)
        for i in range(n_new, n_orig):
            _set_paragraph_text(paragraphs[i], "")
